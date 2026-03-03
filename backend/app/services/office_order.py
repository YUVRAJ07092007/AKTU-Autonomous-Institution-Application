"""Generate Office Order DOCX."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_office_order_docx(
    office_order_no: str,
    institution_name: str,
    application_public_id: str,
    members_list: list[tuple[str, str]],  # (name, role)
    output_path: Path,
) -> None:
    """Write a DOCX file with placeholders filled. Members_list is [(name, role), ...]."""
    doc = DocxDocument()
    doc.add_paragraph("OFFICE ORDER", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph(f"No.: {office_order_no}")
    doc.add_paragraph(f"Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    doc.add_paragraph()
    doc.add_paragraph(f"Institution: {institution_name}")
    doc.add_paragraph(f"Application ID: {application_public_id}")
    doc.add_paragraph()
    doc.add_paragraph("Committee members:")
    for name, role in members_list:
        doc.add_paragraph(f"  - {name} ({role})", style="List Bullet")
    doc.add_paragraph()
    doc.add_paragraph("This committee is constituted for the purpose of Clause 6.29 report.")
    doc.save(str(output_path))
