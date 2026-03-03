"""Generate Meeting Notice DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_meeting_notice_docx(
    application_public_id: str,
    institution_name: str,
    mode: str,
    date_time_str: str,
    venue: str | None,
    online_link: str | None,
    agenda: str,
    output_path: Path,
) -> None:
    """Write meeting notice DOCX."""
    doc = DocxDocument()
    doc.add_paragraph("MEETING NOTICE", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph(f"Application ID: {application_public_id}")
    doc.add_paragraph(f"Institution: {institution_name}")
    doc.add_paragraph()
    doc.add_paragraph(f"Mode: {mode}")
    doc.add_paragraph(f"Date & Time: {date_time_str}")
    if venue:
        doc.add_paragraph(f"Venue: {venue}")
    if online_link:
        doc.add_paragraph(f"Online link: {online_link}")
    doc.add_paragraph()
    doc.add_paragraph("Agenda:")
    doc.add_paragraph(agenda)
    doc.save(str(output_path))
