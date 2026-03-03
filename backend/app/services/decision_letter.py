"""Generate Decision Letter DOCX."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _fmt_date(d: datetime | None) -> str:
    if d is None:
        return "—"
    return d.strftime("%d-%m-%Y")


def generate_decision_letter_docx(
    application_public_id: str,
    institution_name: str,
    decision_type: str,
    tenure_years: int | None,
    valid_from: datetime | None,
    valid_to: datetime | None,
    reasons: str,
    conditions: str,
    ugc_subject_to_flag: bool,
    output_path: Path,
) -> None:
    """Write Decision Letter DOCX and save to output_path."""
    doc = DocxDocument()
    doc.add_paragraph("DECISION LETTER", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Academic Autonomy — AKTU")
    doc.add_paragraph()
    doc.add_paragraph(f"Application ID: {application_public_id}")
    doc.add_paragraph(f"Institution: {institution_name}")
    doc.add_paragraph()
    doc.add_paragraph(f"Decision: {decision_type}", style="Heading 2")
    if tenure_years is not None:
        doc.add_paragraph(f"Tenure: {tenure_years} year(s)")
    doc.add_paragraph(f"Valid from: {_fmt_date(valid_from)}")
    doc.add_paragraph(f"Valid to: {_fmt_date(valid_to)}")
    if ugc_subject_to_flag:
        doc.add_paragraph("Subject to UGC approval.", style="Heading 2")
    if reasons:
        doc.add_paragraph("Reasons:", style="Heading 2")
        doc.add_paragraph(reasons)
    if conditions:
        doc.add_paragraph("Conditions:", style="Heading 2")
        doc.add_paragraph(conditions)
    doc.save(str(output_path))
