"""MoM DOCX generation — Clause 6.29(a)(i), (ii), (iii) style."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Keys used in content_json for structured MoM sections
SECTION_6_29_A_I = "section_6_29_a_i"
SECTION_6_29_A_II = "section_6_29_a_ii"
SECTION_6_29_A_III = "section_6_29_a_iii"
COMMENTS = "comments"

DEFAULT_PLACEHOLDER_I = "[Clause 6.29(a)(i) — Summary of presentation by institution]"
DEFAULT_PLACEHOLDER_II = "[Clause 6.29(a)(ii) — Points raised by committee and response]"
DEFAULT_PLACEHOLDER_III = "[Clause 6.29(a)(iii) — Recommendations / observations]"


def _get_content(content_json: dict | None) -> dict:
    if not content_json:
        return {
            SECTION_6_29_A_I: DEFAULT_PLACEHOLDER_I,
            SECTION_6_29_A_II: DEFAULT_PLACEHOLDER_II,
            SECTION_6_29_A_III: DEFAULT_PLACEHOLDER_III,
            COMMENTS: "",
        }
    return {
        SECTION_6_29_A_I: content_json.get(SECTION_6_29_A_I) or DEFAULT_PLACEHOLDER_I,
        SECTION_6_29_A_II: content_json.get(SECTION_6_29_A_II) or DEFAULT_PLACEHOLDER_II,
        SECTION_6_29_A_III: content_json.get(SECTION_6_29_A_III) or DEFAULT_PLACEHOLDER_III,
        COMMENTS: content_json.get(COMMENTS) or "",
    }


def generate_mom_draft_docx(
    application_public_id: str,
    institution_name: str,
    content_json: dict | None,
    output_path: Path,
) -> None:
    """Write MoM draft DOCX with placeholders or existing content (Clause 6.29(a)/(d) style)."""
    content = _get_content(content_json)
    doc = DocxDocument()
    doc.add_paragraph("MINUTES OF MEETING", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("(As per AKTU Clause 6.29(a) / 6.29(d))")
    doc.add_paragraph()
    doc.add_paragraph(f"Application ID: {application_public_id}")
    doc.add_paragraph(f"Institution: {institution_name}")
    doc.add_paragraph()
    doc.add_paragraph("6.29(a)(i) — Summary of presentation by the institution:", style="Heading 2")
    doc.add_paragraph(content[SECTION_6_29_A_I])
    doc.add_paragraph()
    doc.add_paragraph("6.29(a)(ii) — Points raised by the committee and response of the institution:", style="Heading 2")
    doc.add_paragraph(content[SECTION_6_29_A_II])
    doc.add_paragraph()
    doc.add_paragraph("6.29(a)(iii) — Recommendations / observations:", style="Heading 2")
    doc.add_paragraph(content[SECTION_6_29_A_III])
    if content[COMMENTS]:
        doc.add_paragraph()
        doc.add_paragraph("Comments / change log:", style="Heading 2")
        doc.add_paragraph(content[COMMENTS])
    doc.save(str(output_path))


def render_mom_final_docx(
    application_public_id: str,
    institution_name: str,
    content_json: dict,
    output_path: Path,
) -> None:
    """Render final MoM DOCX from structured content (no placeholders)."""
    generate_mom_draft_docx(
        application_public_id=application_public_id,
        institution_name=institution_name,
        content_json=content_json,
        output_path=output_path,
    )
